import os
from typing import Dict, Any

class DOCXParser:
    """Parser using python-docx to extract text from Word documents."""

    def parse(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Extract table texts too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            metadata = {
                "title": os.path.basename(file_path),
                "format": "DOCX"
            }
            return {
                "text": "\n".join(full_text),
                "metadata": metadata
            }

        except ImportError:
            return {
                "text": f"[DOCX Library missing: python-docx. Unable to parse file {file_path}]",
                "metadata": {"title": os.path.basename(file_path)}
            }
        except Exception as e:
            raise RuntimeError(f"Error parsing Word file: {e}")
